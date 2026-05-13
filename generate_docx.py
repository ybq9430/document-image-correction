"""Generate Technical Report as .docx with embedded images."""
import os, sys
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.font_manager as fm

# Set Korean font
font_path = None
for f in fm.fontManager.ttflist:
    if 'Malgun' in f.name or 'Gulim' in f.name:
        font_path = f.fname
        break
if font_path:
    fm.fontManager.addfont(font_path)
    prop = fm.FontProperties(fname=font_path)
    plt.rcParams['font.family'] = prop.get_name()
    print(f"Using font: {prop.get_name()} ({font_path})")
else:
    print("WARNING: No Korean font found, Hangul will show as squares")
import numpy as np
import cv2

OUTPUT_DIR = r"D:\claude02\项目\document_correction\output"
STEPS_DIR_TEMPLATE = os.path.join(OUTPUT_DIR, "{}_steps")
DOCX_PATH = r"D:\claude02\项目\Technical_Report.docx"

# ── Helper functions ──
def add_heading_kr(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'Malgun Gothic'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')
    return h

def add_para(doc, text, bold=False, size=10, align=None, space_after=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Malgun Gothic'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')
    run.font.size = Pt(size)
    run.bold = bold
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    return p

def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri+1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
    doc.add_paragraph()
    return table

def read_img_safe(fpath):
    """Read image with Unicode path support for Windows."""
    try:
        img = cv2.imread(fpath)
        if img is not None:
            return img
    except:
        pass
    # Fallback: use numpy from file
    try:
        with open(fpath, 'rb') as f:
            data = np.frombuffer(f.read(), dtype=np.uint8)
        return cv2.imdecode(data, cv2.IMREAD_COLOR)
    except:
        return None

def add_image(doc, img_path, width_inches=5.5, caption=None):
    if not os.path.exists(img_path):
        add_para(doc, f"[Image missing / 이미지 없음: {os.path.basename(img_path)}]", size=9)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(img_path, width=Inches(width_inches))
    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cap.add_run(caption)
        r.font.size = Pt(9)
        r.italic = True

def generate_step_grid(img_name, save_path):
    """Create a grid of 7 pipeline steps for one image."""
    step_files = [
        ("01_original.jpg", "① 원본"),
        ("02_grayscale.jpg", "② Grayscale"),
        ("03_blurred.jpg", "③ Gaussian Blur"),
        ("04_edges.jpg", "④ Canny Edges"),
        ("05_contours.jpg", "⑤ Contours"),
        ("06_perspective.jpg", "⑥ Perspective"),
        ("07_enhanced.jpg", "⑦ Enhanced"),
    ]
    steps_dir = STEPS_DIR_TEMPLATE.format(img_name)
    fig, axes = plt.subplots(2, 4, figsize=(16, 8), dpi=150)
    axes = axes.flatten()
    for i, (fname, title) in enumerate(step_files):
        fpath = os.path.join(steps_dir, fname)
        if os.path.exists(fpath):
            img = read_img_safe(fpath)
            if img is not None:
                img_rgb = cv2.cvtColor(img, cv2.COLOR_BGR2RGB) if len(img.shape)==3 else img
                axes[i].imshow(img_rgb, cmap='gray' if len(img.shape)==2 else None)
        axes[i].set_title(title, fontsize=10, fontweight='bold')
        axes[i].axis('off')
    axes[7].axis('off')  # hide 8th panel
    fig.suptitle(f"Pipeline Steps / 파이프라인 단계 — {img_name}", fontsize=14, fontweight='bold')
    plt.tight_layout()
    fig.savefig(save_path, dpi=150, bbox_inches='tight', facecolor='white')
    plt.close(fig)

def generate_comparison_figure(img_name, save_path):
    """Generate 3-panel comparison: Original | Perspective | Enhanced."""
    steps_dir = STEPS_DIR_TEMPLATE.format(img_name)
    orig_path = os.path.join(steps_dir, "01_original.jpg")
    persp_path = os.path.join(steps_dir, "06_perspective.jpg")
    enh_path = os.path.join(steps_dir, "07_enhanced.jpg")

    fig, axes = plt.subplots(1, 3, figsize=(15, 5), dpi=150)
    titles = ["Original / 원본", "Perspective Corrected / 투시 보정", "Enhanced / 향상 결과"]
    paths = [orig_path, persp_path, enh_path]

    for ax, title, fpath in zip(axes, titles, paths):
        if os.path.exists(fpath):
            img = read_img_safe(fpath)
            if img is not None:
                display_img = cv2.cvtColor(img, cv2.COLOR_BGR2RGB) if len(img.shape)==3 else img
                ax.imshow(display_img, cmap='gray' if len(img.shape)==2 else None)
        ax.set_title(title, fontsize=11, fontweight='bold')
        ax.axis('off')

    fig.suptitle(f"Document Correction Result — {img_name}", fontsize=14, fontweight='bold')
    plt.tight_layout()
    fig.savefig(save_path, dpi=150, bbox_inches='tight', facecolor='white')
    plt.close(fig)

def generate_sensitivity_chart(save_path):
    """Generate parameter sensitivity analysis chart from experimental data."""
    fig, axes = plt.subplots(2, 2, figsize=(14, 10), dpi=150)

    # Experiment 1: Canny Low
    ax = axes[0, 0]
    lows = [50, 75, 100, 125]
    psnr1 = [5.1, 5.1, 5.1, 5.1]
    ssim1 = [0.4516, 0.4516, 0.4543, 0.4564]
    lap1 = [130.4, 130.4, 127.5, 125.9]
    ax2 = ax.twinx()
    ax.plot(lows, psnr1, 'b-o', label='PSNR (dB)')
    ax.plot(lows, lap1, 'g-s', label='Laplacian Ratio')
    ax2.plot(lows, [s*100 for s in ssim1], 'r-^', label='SSIM x100')
    ax.set_title('Canny Low Threshold', fontweight='bold')
    ax.set_xlabel('Low Threshold'); ax.set_ylabel('PSNR / Laplacian Ratio')
    ax2.set_ylabel('SSIM x100'); ax.grid(True, alpha=0.3)
    lines1 = ax.get_lines() + ax2.get_lines()
    ax.legend(lines1, [l.get_label() for l in lines1], fontsize=8)

    # Experiment 2: Canny High
    ax = axes[0, 1]
    highs = [150, 200, 250, 300]
    psnr2 = [5.1, 5.1, 5.1, 4.9]
    ssim2 = [0.4516, 0.4516, 0.4516, 0.0]
    lap2 = [130.4, 130.4, 130.4, 0.0]
    ax2 = ax.twinx()
    ax.plot(highs, psnr2, 'b-o', label='PSNR (dB)')
    ax.plot(highs, lap2, 'g-s', label='Laplacian Ratio')
    ax2.plot(highs, [s*100 for s in ssim2], 'r-^', label='SSIM x100')
    ax.set_title('Canny High Threshold', fontweight='bold')
    ax.set_xlabel('High Threshold'); ax.set_ylabel('PSNR / Laplacian Ratio')
    ax2.set_ylabel('SSIM x100'); ax.grid(True, alpha=0.3)
    ax.axvline(x=300, color='red', linestyle='--', alpha=0.3, linewidth=3, label='Detection failure')
    lines2 = ax.get_lines() + ax2.get_lines()
    ax.legend(lines2, [l.get_label() for l in lines2], fontsize=8)

    # Experiment 3: Gaussian Kernel
    ax = axes[1, 0]
    kernels = ['3x3', '5x5', '7x7', '9x9']
    psnr3 = [5.1, 5.1, 5.1, 4.9]
    ssim3 = [0.4511, 0.4516, 0.4511, 0.6399]
    lap3 = [130.8, 130.4, 130.8, 2.7]
    ax2 = ax.twinx()
    ax.plot(range(4), psnr3, 'b-o', label='PSNR (dB)')
    ax.plot(range(4), lap3, 'g-s', label='Laplacian Ratio')
    ax2.plot(range(4), [s*100 for s in ssim3], 'r-^', label='SSIM x100')
    ax.set_title('Gaussian Kernel Size', fontweight='bold')
    ax.set_xticks(range(4)); ax.set_xticklabels(kernels)
    ax.set_xlabel('Kernel Size'); ax.set_ylabel('PSNR / Laplacian Ratio')
    ax2.set_ylabel('SSIM x100'); ax.grid(True, alpha=0.3)
    ax.axvline(x=3, color='red', linestyle='--', alpha=0.3, linewidth=3)
    lines3 = ax.get_lines() + ax2.get_lines()
    ax.legend(lines3, [l.get_label() for l in lines3], fontsize=8)

    # Experiment 4: CLAHE clipLimit
    ax = axes[1, 1]
    clips = [1.0, 2.0, 3.0, 4.0, 5.0]
    psnr4 = [5.0, 5.1, 5.1, 5.1, 5.2]
    ssim4 = [0.5016, 0.4516, 0.4177, 0.3882, 0.3656]
    lap4 = [94.4, 130.4, 156.7, 178.5, 195.9]
    ax2 = ax.twinx()
    ax.plot(clips, psnr4, 'b-o', label='PSNR (dB)')
    ax.plot(clips, lap4, 'g-s', label='Laplacian Ratio')
    ax2.plot(clips, [s*100 for s in ssim4], 'r-^', label='SSIM x100')
    ax.set_title('CLAHE clipLimit', fontweight='bold')
    ax.set_xlabel('clipLimit'); ax.set_ylabel('PSNR / Laplacian Ratio')
    ax2.set_ylabel('SSIM x100'); ax.grid(True, alpha=0.3)
    ax.axvline(x=2.0, color='green', linestyle='--', alpha=0.5, linewidth=2, label='Optimal')
    lines4 = ax.get_lines() + ax2.get_lines()
    ax.legend(lines4, [l.get_label() for l in lines4], fontsize=8)

    fig.suptitle('Parameter Sensitivity Analysis / 파라미터 민감도 분석', fontsize=14, fontweight='bold')
    plt.tight_layout()
    fig.savefig(save_path, dpi=150, bbox_inches='tight', facecolor='white')
    plt.close(fig)

def generate_flowchart(save_path):
    """Generate algorithm pipeline flowchart."""
    fig, ax = plt.subplots(1, 1, figsize=(12, 6), dpi=150)
    ax.set_xlim(0, 12)
    ax.set_ylim(0, 10)
    ax.axis('off')

    steps = [
        (1, 9, "Input Image\n입력 이미지", "#E3F2FD"),
        (3, 9, "① Grayscale\n그레이스케일", "#FFF3E0"),
        (5, 9, "② Gaussian Blur\n가우시안 블러", "#FFF3E0"),
        (7, 9, "③ Canny Edge\n에지 검출", "#E8F5E9"),
        (9, 9, "④ Contour + Hull\n윤곽선 검출", "#E8F5E9"),
        (9, 7, "⑤ Perspective\n투시 변환", "#F3E5F5"),
        (7, 5, "⑥ CLAHE + Threshold\n적응형 이진화", "#FCE4EC"),
        (5, 5, "⑦ Morphology\n형태학 연산", "#FCE4EC"),
        (3, 5, "Output\n출력 이미지", "#E3F2FD"),
    ]

    for x, y, label, color in steps:
        ax.add_patch(plt.Rectangle((x-0.8, y-0.5), 1.6, 1.0,
                                    facecolor=color, edgecolor='black', linewidth=1.5))
        ax.text(x, y, label, ha='center', va='center', fontsize=8, fontweight='bold')

    # Arrows
    arrows = [(1.8,9, 2.2,9), (3.8,9, 4.2,9), (5.8,9, 6.2,9), (7.8,9, 8.2,9),
              (9.8,8.5, 9.8,7.5), (8.2,7, 7.8,5.5), (6.2,5.5, 5.8,5), (4.2,5, 3.8,5)]
    for x1, y1, x2, y2 in arrows:
        ax.annotate('', xy=(x2, y2), xytext=(x1, y1),
                    arrowprops=dict(arrowstyle='->', color='gray', lw=1.5))

    # Module labels
    ax.text(2, 9.8, "[MODULE 1: 전처리]", fontsize=7, color='gray', ha='center')
    ax.text(8, 9.8, "[MODULE 2: 문서 검출]", fontsize=7, color='gray', ha='center')
    ax.text(9, 6.2, "[MODULE 3: 기하보정]", fontsize=7, color='gray', ha='center', rotation=90)
    ax.text(4, 5.8, "[MODULE 4: 화질 향상]", fontsize=7, color='gray', ha='center')

    ax.set_title('Document Correction Pipeline / 문서 교정 파이프라인', fontsize=14, fontweight='bold')
    fig.savefig(save_path, dpi=150, bbox_inches='tight', facecolor='white')
    plt.close(fig)

# ============================================================
# BUILD DOCUMENT
# ============================================================
print("Generating DOCX...")
doc = Document()

# Page setup
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# ── Title Page ──
add_para(doc, "", size=10)
add_para(doc, "Technical Report (TR)", bold=True, size=22, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para(doc, "전통적 영상처리 기법 기반", bold=True, size=14, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para(doc, "문서 이미지 자동 교정 및 향상 시스템 설계", bold=True, size=14, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para(doc, "", size=10)
add_para(doc, "과목명: 고급컴퓨터비전 (2026학년도 1학기)", size=11, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para(doc, "제출일: 2026년 5월", size=11, align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_page_break()

# ============================================================
# 제1장 TR 개요
# ============================================================
add_heading_kr(doc, "제1장. TR 개요", level=1)

add_heading_kr(doc, "1.1 연구 배경 및 의의", level=2)
add_para(doc, "최근 디지털 전환이 가속화됨에 따라 종이 문서를 디지털 형태로 변환하는 수요가 급격히 증가하고 있다. 기업, 정부 기관, 교육 기관 등에서는 대량의 문서를 효율적으로 관리하기 위해 문서 스캔 및 OCR(Optical Character Recognition) 기술을 적극적으로 활용하고 있다. 특히 스마트폰의 보급으로 인해 전용 스캐너 없이도 카메라를 이용한 문서 촬영이 일상화되었으나, 이러한 방식으로 획득된 문서 이미지에는 다양한 품질 문제가 발생한다.")
add_para(doc, "스마트폰으로 문서를 촬영할 때 발생하는 주요 문제점은 다음과 같다. 첫째, 촬영 각도에 의한 투시 왜곡(perspective distortion)으로 문서가 사다리꼴 형태로 변형된다. 둘째, 실내 조명이나 자연광에 의한 불균일한 조명(uneven illumination)이 발생하여 문서의 일부 영역이 과도하게 밝거나 어두워진다. 셋째, 카메라 센서의 한계로 인한 노이즈(noise)가 이미지 품질을 저하시킨다. 넷째, 낮은 대비(contrast)로 인해 텍스트의 가독성이 떨어지는 문제가 있다.")
add_para(doc, "이러한 문제를 해결하기 위해 딥러닝 기반 접근법이 연구되고 있으나, 높은 계산 비용과 대규모 학습 데이터의 필요성, 그리고 모델의 해석 불가능성 등의 한계가 존재한다. 반면, 전통적 영상처리 기법은 수학적 원리에 기반하여 결과의 해석이 용이하고, 계산 비용이 낮으며, GPU 없이도 실시간 처리가 가능하다는 장점을 갖는다. 따라서 본 프로젝트에서는 전통적 컴퓨터 비전 및 영상처리 기법만을 활용하여 문서 이미지의 자동 교정 및 향상 시스템을 설계하고 구현하였다.")

add_heading_kr(doc, "1.2 연구 가설", level=2)
add_para(doc, "본 프로젝트의 연구 가설은 다음과 같다: 문서 이미지에서 문서의 외곽 경계가 충분히 검출 가능한 경우, Canny 에지 검출, 윤곽선 근사, 투시 변환, CLAHE, 적응적 이진화와 같은 전통적 영상처리 기법만으로도 문서의 기하학적 왜곡을 효과적으로 보정하고 텍스트 가독성을 향상시킬 수 있다.")
add_para(doc, "즉, 스마트폰 촬영 과정에서 기울어지거나 사다리꼴 형태의 원근 왜곡이 발생하더라도, 문서 외곽을 안정적으로 찾을 수 있다면 딥러닝 기반 모델 없이도 충분히 실용적인 수준의 문서 교정 결과를 얻을 수 있다는 가정이다. 또한 조명 불균일, 저대비, 미세 잡음과 같은 품질 저하 요소에 대해서도 CLAHE, Adaptive Thresholding, Morphological Processing의 조합이 가독성 향상에 유의미한 효과를 줄 것으로 예상한다.")

add_heading_kr(doc, "1.3 연구 목표", level=2)
add_para(doc, "본 프로젝트의 핵심 목표는 스마트폰으로 촬영된 문서 이미지를 입력으로 받아 기하학적 교정과 화질 향상을 자동으로 수행하는 완전한 처리 파이프라인을 구축하는 것이다. 구체적인 연구 목표는: (1) 자동 문서 영역 검출 — 복잡한 배경에서 문서 영역을 자동으로 식별하고 경계를 추출, (2) 투시 기하 교정 — 검출된 문서의 네 꼭짓점을 기반으로 투시 변환을 적용하여 직사각형 형태로 교정, (3) 이미지 품질 향상 — 적응적 이진화, 대비 향상, 형태학적 연산을 통해 문서의 가독성 향상, (4) 완전한 처리 흐름 구축 — 딥러닝을 사용하지 않고 전통적 영상처리 기법만으로 입력부터 출력까지의 일관된 파이프라인 구현이다.")

add_heading_kr(doc, "1.4 개발 환경 및 제약 조건", level=2)
add_table(doc, ["항목", "내용"], [
    ["프로그래밍 언어", "Python 3.x"],
    ["핵심 라이브러리", "OpenCV 4.13.0, NumPy 2.0.2, Matplotlib, scikit-image, pandas"],
    ["개발 환경", "Google Colab (CPU 런타임)"],
    ["ML/DL 사용", "금지 (torch, tensorflow, keras, sklearn 등 불가)"],
    ["허용 기법", "전통적 Computer Vision + Image Processing만 사용"],
])

add_heading_kr(doc, "1.5 시스템 전체 흐름도", level=2)
# Generate flowchart
flowchart_path = os.path.join(OUTPUT_DIR, "_flowchart.png")
generate_flowchart(flowchart_path)
add_image(doc, flowchart_path, width_inches=5.5, caption="그림 1. 문서 이미지 자동 교정 및 향상 시스템 알고리즘 흐름도")

doc.add_page_break()

# ============================================================
# 제2장 TR 내용
# ============================================================
add_heading_kr(doc, "제2장. TR 내용", level=1)

add_heading_kr(doc, "2.1 전체 알고리즘 구조", level=2)
add_para(doc, "전체 알고리즘의 처리 파이프라인은 11단계로 구성된다: (1) Input Document Image → (2) Grayscale Conversion → (3) Gaussian Filtering → (4) Canny Edge Detection → (4b) Edge Dilation → (5) Contour Detection → (6) Quadrilateral Approximation + Convex Hull → (7) Corner Point Ordering → (8) Perspective Transformation → (9) CLAHE / Adaptive Thresholding → (10) Morphological Processing → (11) Output Corrected Image.")
add_para(doc, "이 파이프라인은 기능별로 네 개의 모듈로 구분된다: ① 전처리(Preprocessing) — Grayscale 변환, Gaussian Filtering, ② 문서 검출(Detection) — Canny Edge, Edge Dilation, Contour, Quadrilateral Approximation + Convex Hull, ③ 기하학적 보정(Geometry) — Corner Ordering, Perspective Transform, ④ 화질 향상(Enhancement) — CLAHE, Adaptive Thresholding, Morphological Processing.")

add_heading_kr(doc, "2.2 단계별 알고리즘 상세", level=2)

add_para(doc, "Step 2: Grayscale Conversion — 입력 RGB 이미지를 Grayscale로 변환하여 채널 수를 3에서 1로 줄임으로써 후속 처리의 계산 효율성을 높인다. 공식은 Gray = 0.299×R + 0.587×G + 0.114×B 이며, cv2.cvtColor(img, cv2.COLOR_BGR2GRAY) 함수를 사용한다.", bold=False)
add_para(doc, "Step 3: Gaussian Filtering — 5×5 커널, 시그마 0(자동계산)의 가우시안 필터를 적용하여 고주파 노이즈를 제거하고 영상을 평활화한다. cv2.GaussianBlur(gray, (5,5), 0) 함수를 사용한다.")
add_para(doc, "Step 4: Canny Edge Detection — 문서의 외곽 경계 후보를 추출한다. 하한 임계값 75, 상한 임계값 200을 사용하며, 내부적으로 Gaussian smoothing → Gradient calculation (Sobel) → Non-maximum Suppression → Double Thresholding → Hysteresis Edge Tracking 순서로 동작한다.")
add_para(doc, "Step 4b: Edge Dilation — Canny 에지 검출 결과의 조각난 에지 세그먼트를 연결하기 위해 3×3 Morphological Rect 커널로 2회 반복 팽창(dilation)을 수행한다. cv2.dilate() 함수를 사용한다.")
add_para(doc, "Step 5-6: Contour Detection + Quadrilateral Approximation — 본 시스템의 핵심인 3단계 적응형 문서 검출 전략: (1) Direct Approximation — 각 contour에 Douglas-Peucker 알고리즘(cv2.approxPolyDP)을 적용하고 epsilon 비율을 0.02~0.10 범위에서 자동 탐색, (2) Convex Hull Approximation — 개별 contour의 convex hull을 계산 후 재근사, (3) Combined Convex Hull Fallback — 모든 contour 결합 후 convex hull 근사. 최종적으로 면적이 전체 이미지의 2% 미만인 사각형은 노이즈로 판단하여 제외한다.")
add_para(doc, "Step 7-8: Corner Ordering + Perspective Transform — 검출된 4개 꼭짓점을 TL(min x+y), TR(max x-y), BR(max x+y), BL(min x-y) 순서로 정렬한 후, cv2.getPerspectiveTransform()으로 3×3 Homography Matrix를 계산하고 cv2.warpPerspective()로 투시 변환을 적용한다. 출력 크기는 width = max(dist(BL,BR), dist(TL,TR)), height = max(dist(TR,BR), dist(TL,BL))로 자동 계산된다.")
add_para(doc, "Step 9-10: Enhancement — CLAHE(clipLimit=2.0, tileGridSize=8×8)로 지역 대비를 향상시키고, Adaptive Thresholding(Gaussian method, blockSize=11, C=2)으로 조명 불균일을 보정한다. 이후 Morphological Close(3×3 kernel)로 문자 내부 구멍을 메우고 배경 잡음을 제거한다.")

add_heading_kr(doc, "2.3 파이프라인 처리 결과 (중간 단계)", level=2)
add_para(doc, "다음 그림은 img01에 대한 7단계 파이프라인 중간 결과를 보여준다.")

step_grid_path = os.path.join(OUTPUT_DIR, "_step_grid.png")
generate_step_grid("img01", step_grid_path)
add_image(doc, step_grid_path, width_inches=5.5, caption="그림 2. img01 파이프라인 7단계 중간 결과")

add_heading_kr(doc, "2.4 실험 설계", level=2)
add_para(doc, "실험에는 Samsung Galaxy S24 Ultra로 직접 촬영한 6장의 문서 이미지(1280×1706)를 사용하였다. 촬영 조건은 약 15°~45° 기울기, 실내 형광등 조명, 책상 위 A4 문서이다.")

add_heading_kr(doc, "2.5 평가 지표", level=2)
add_table(doc, ["#", "지표", "측정 방법"], [
    ["①", "문서 영역 검출 성공률", "성공 이미지 수 / 전체 이미지 수 × 100%"],
    ["②", "처리 시간", "time.time() 측정 (초), mean/min/max/total"],
    ["③", "대비 향상도", "PSNR(dB), Laplacian Variance(전/후/배율), SSIM(0~1)"],
    ["④", "처리 전후 시각 비교", "원본 / 투시 교정 / 최종 향상 3단 비교도"],
])

add_heading_kr(doc, "2.6 파라미터 설정", level=2)
add_table(doc, ["파라미터", "기본값", "설명"], [
    ["Gaussian kernel", "(5, 5)", "노이즈 제거 강도"],
    ["Canny low", "75", "약한 에지 임계값"],
    ["Canny high", "200", "강한 에지 임계값"],
    ["approxPolyDP epsilon", "0.02~0.10 (자동 탐색)", "다각형 근사 정밀도"],
    ["CLAHE clipLimit", "2.0", "대비 향상 강도"],
    ["CLAHE tileGridSize", "(8, 8)", "지역 처리 크기"],
    ["Adaptive blockSize", "11", "지역 이진화 블록 크기"],
    ["Adaptive C", "2", "이진화 보정값"],
    ["Morph kernel", "(3, 3)", "형태학적 연산 크기"],
])

doc.add_page_break()

# ============================================================
# 제3장 TR 고찰 및 분석
# ============================================================
add_heading_kr(doc, "제3장. TR 고찰 및 분석", level=1)

add_heading_kr(doc, "3.1 실험 결과", level=2)

add_para(doc, "6장의 테스트 이미지에 대한 처리 결과는 다음 표와 같다. 모든 실험은 Google Colab CPU 환경에서 수행되었으며, 각 이미지당 3회 반복 측정의 평균값을 기록하였다.", bold=False)

add_table(doc,
    ["이미지명", "문서 검출", "처리 시간(s)", "PSNR(dB)", "Laplacian 전", "Laplacian 후", "Laplacian 배율", "SSIM"],
    [
        ["img01", "성공", "0.126", "5.07", "77.33", "10,083.79", "130.40x", "0.4516"],
        ["img02", "성공", "0.096", "6.59", "26.81", "8,213.48", "306.41x", "0.5938"],
        ["img03", "성공", "0.120", "6.50", "67.67", "10,916.22", "161.31x", "0.5213"],
        ["img04", "성공", "0.097", "6.09", "80.07", "10,441.34", "130.40x", "0.5173"],
        ["img05", "성공", "0.098", "6.32", "68.96", "9,906.58", "143.65x", "0.5411"],
        ["img06", "성공", "0.111", "6.57", "74.31", "9,479.85", "127.56x", "0.5375"],
        ["평균", "100%", "0.108", "6.19", "65.86", "9,840.21", "166.62x", "0.5271"],
    ])

add_heading_kr(doc, "3.2 결과 분석", level=2)

add_para(doc, "① 문서 영역 검출 성공률: 6/6 = 100.0%", bold=True)
add_para(doc, "3단계 적응형 검출 전략(직접 근사 → Convex Hull → 결합 Hull)을 통해 모든 테스트 이미지에서 문서 영역을 성공적으로 검출하였다. 특히 문서가 프레임을 대부분 채우는 경우에도 결합 Convex Hull 전략이 효과적으로 작동하여 100% 검출률을 달성할 수 있었다.")

add_para(doc, "② 처리 시간: 평균 0.108초 (최소 0.096초, 최대 0.126초)", bold=True)
add_para(doc, "Google Colab CPU 환경에서 단일 이미지(1280×1706) 처리에 평균 0.108초가 소요되었다. 이는 NFR-02에서 정의한 '단일 이미지 처리 시간 5초 이내' 요구사항을 크게 상회하는 결과로, 실시간 처리도 충분히 가능한 수준이다. 전체 6장 처리 시간은 약 0.648초에 불과하다.")

add_para(doc, "③ 대비 향상도", bold=True)
add_para(doc, "Laplacian Variance: 평균 65.86 → 9,840.21 (약 166.62배 향상). Laplacian 분산값은 이미지의 선명도(sharpness)를 나타내는 지표로, 값이 높을수록 더 선명한 이미지를 의미한다. 처리 전 평균 65.86에서 처리 후 평균 9,840.21로 약 166.62배 증가하여, 적응적 이진화와 CLAHE를 통한 화질 향상이 매우 효과적이었음을 확인할 수 있다.")
add_para(doc, "PSNR: 평균 6.19 dB — PSNR 값이 상대적으로 낮은 이유는 Adaptive Thresholding을 통한 이진화 과정에서 원본 grayscale 이미지가 흑백 이진 이미지로 변환되면서 픽셀 값의 큰 변화가 발생했기 때문이다. 오히려 SSIM과 Laplacian Ratio가 문서 교정 품질 평가에 더 적합한 지표이다.")
add_para(doc, "SSIM: 평균 0.5271 (0.4516~0.5938 범위) — 이진화 처리의 특성상 원본 grayscale 이미지와 이진화된 결과 간의 구조적 차이를 반영한다. SSIM이 '원본과의 유사성'을 측정하는 반면, 본 시스템의 목표는 '가독성 향상'이기 때문에 SSIM만으로 시스템의 효과를 판단할 수 없음을 시사한다.")

# Add comparison figure
comp_path = os.path.join(OUTPUT_DIR, "_comparison.png")
generate_comparison_figure("img01", comp_path)
add_image(doc, comp_path, width_inches=5.5, caption="그림 3. img01 처리 전후 비교 (원본 → 투시 교정 → 최종 향상)")

add_heading_kr(doc, "3.3 파라미터 민감도 실험", level=2)
add_para(doc, "대표 이미지(img01)에 대해 4가지 주요 파라미터(Canny low, Canny high, Gaussian kernel, CLAHE clipLimit)의 영향을 실험하였다. 실험 결과는 아래 차트와 같다.")

# Add sensitivity chart
sens_path = os.path.join(OUTPUT_DIR, "_sensitivity.png")
generate_sensitivity_chart(sens_path)
add_image(doc, sens_path, width_inches=5.5, caption="그림 4. 파라미터 민감도 분석 — 4가지 파라미터가 PSNR, SSIM, Laplacian Ratio에 미치는 영향")

add_para(doc, "Canny Low Threshold (50~125): 모든 값에서 검출 성공, PSNR/SSIM/Laplacian Ratio 모두 유사. Edge Dilation이 작은 변화를 흡수하기 때문.", bold=False)
add_para(doc, "Canny High Threshold (150~300): high=300에서 SSIM NaN, Laplacian Ratio 0 — 임계값 과대로 인한 에지 소실. 150~250 범위가 안정적.", bold=False)
add_para(doc, "Gaussian Kernel (3×3~9×9): (9,9)에서 Laplacian Ratio 2.7x로 급감 — 과도한 블러링으로 문서 경계 부정확. (5,5)가 최적.", bold=False)
add_para(doc, "CLAHE clipLimit (1.0~5.0): clipLimit 증가 시 Laplacian Ratio 94.4x→195.9x 상승, SSIM 0.5016→0.3656 하락. 선명도와 구조 보존 간 균형을 고려할 때 clipLimit=2.0이 최적.", bold=False)

add_heading_kr(doc, "3.4 종합 파라미터 권장값", level=2)
add_table(doc, ["파라미터", "권장값", "근거"], [
    ["Canny low", "75", "넓은 범위(50~125)에서 안정적"],
    ["Canny high", "200", "150~250에서 안정적, 300에서 실패"],
    ["Gaussian kernel", "(5,5)", "(3,3)~(7,7)에서 안정적, (9,9)에서 성능 저하"],
    ["CLAHE clipLimit", "2.0", "Laplacian 130.4x + SSIM 0.4516 최적 균형"],
])

add_heading_kr(doc, "3.5 가설 검증", level=2)
add_para(doc, "본 연구의 가설은 '전통적 영상처리 기법만으로도 문서의 기하학적 왜곡을 효과적으로 보정하고 텍스트 가독성을 향상시킬 수 있다'는 것이었다. 실험 결과, 100%의 문서 검출 성공률, 평균 166.62배의 Laplacian 선명도 향상, 평균 0.108초의 처리 속도를 달성하여 가설이 지지됨을 확인하였다.")

add_heading_kr(doc, "3.6 한계점 및 향후 연구", level=2)
add_para(doc, "1. 복잡한 배경: 본 실험 이미지는 상대적으로 단순한 배경(책상)에서 촬영되었다. 매우 복잡한 배경에서는 문서 검출 성능이 저하될 수 있어, 텍스처 기반 배경 제거나 적응형 ROI 탐색 기법을 추가할 수 있다.")
add_para(doc, "2. PSNR의 낮은 값: 이진화로 인한 PSNR 저하는 불가피하다. OCR 인식률과 같은 실제 활용 지표를 추가 평가 지표로 도입하는 것이 바람직하다.")
add_para(doc, "3. Grayscale 보존 모드: 문서 유형에 따라 이진화 대신 Grayscale 보존 모드를 선택적으로 적용하는 하이브리드 접근법을 고려할 수 있다.")
add_para(doc, "4. 파라미터 자동 보정: 이미지 특성(조도, 대비, 해상도)에 따라 파라미터를 자동으로 조정하는 적응형 메커니즘을 추가하면 더 넓은 범위의 입력에 대응할 수 있을 것이다.")

add_heading_kr(doc, "3.7 결론", level=2)
add_para(doc, "본 프로젝트에서는 전통적 컴퓨터 비전 및 영상처리 기법만을 사용하여 스마트폰 촬영 문서 이미지의 자동 교정 및 향상 시스템을 성공적으로 설계하고 구현하였다. Canny 에지 검출, Convex Hull 기반 적응형 문서 검출, 투시 변환, CLAHE, 적응적 이진화, 형태학적 연산을 조합한 11단계 파이프라인을 구축하였다.")
add_para(doc, "6장의 테스트 이미지에 대한 실험 결과, 100% 문서 검출 성공률, 평균 0.108초의 처리 속도, 평균 166.62배의 Laplacian 선명도 향상을 달성하였다. 파라미터 민감도 실험을 통해 최적의 파라미터 조합을 도출하였으며, 모든 평가 지표에서 안정적인 성능을 확인하였다.")
add_para(doc, "본 연구는 딥러닝을 사용하지 않고도 전통적 영상처리 기법만으로 실용적인 문서 교정 시스템을 구축할 수 있음을 입증하였으며, 향후 OCR 연계, 모바일 최적화, 실시간 비디오 처리 등으로 확장할 수 있는 기반을 마련하였다.")

doc.add_page_break()

# ── References ──
add_heading_kr(doc, "참고 문헌", level=1)
refs = [
    "[1] J. Canny, \"A Computational Approach to Edge Detection,\" IEEE TPAMI, vol. 8, no. 6, pp. 679-698, 1986.",
    "[2] R. C. Gonzalez and R. E. Woods, Digital Image Processing, 4th ed., Pearson, 2018.",
    "[3] G. Bradski and A. Kaehler, Learning OpenCV, O'Reilly Media, 2008.",
    "[4] R. Hartley and A. Zisserman, Multiple View Geometry in Computer Vision, 2nd ed., Cambridge University Press, 2004.",
    "[5] K. Zuiderveld, \"Contrast Limited Adaptive Histogram Equalization,\" in Graphics Gems IV, pp. 474-485, 1994.",
    "[6] D. H. Douglas and T. K. Peucker, \"Algorithms for the Reduction of the Number of Points Required to Represent a Digitized Line,\" Cartographica, vol. 10, no. 2, pp. 112-122, 1973.",
    "[7] S. Suzuki and K. Abe, \"Topological Structural Analysis of Digitized Binary Images by Border Following,\" CVGIP, vol. 30, no. 1, pp. 32-46, 1985.",
    "[8] J. Sauvola and M. Pietikäinen, \"Adaptive Document Image Binarization,\" Pattern Recognition, vol. 33, no. 2, pp. 225-236, 2000.",
]
for ref in refs:
    add_para(doc, ref, size=9, space_after=3)

# ── Save ──
doc.save(DOCX_PATH)
print(f"DOCX saved to: {DOCX_PATH}")
print("Done!")
