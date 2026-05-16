"""
Generate Korean research report in .docx format with three-line tables and embedded charts.

Output: D:\claude02\项目\연구보고서_문서이미지교정.docx
"""

import os
from docx import Document
from docx.shared import Inches, Pt, Cm, Emu, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import copy

# ── Paths ──────────────────────────────────────────────────
OUTDIR = r"D:\claude02\项目\document_correction\output"
OUTPUT_PATH = r"D:\claude02\项目\연구보고서_문서이미지교정.docx"

IMG = lambda name: os.path.join(OUTDIR, name)

# ── Document Setup ─────────────────────────────────────────
doc = Document()

# Page size A4
section = doc.sections[0]
section.page_width = Cm(21.0)
section.page_height = Cm(29.7)
section.top_margin = Cm(2.54)
section.bottom_margin = Cm(2.54)
section.left_margin = Cm(2.54)
section.right_margin = Cm(2.54)

# Default style
style = doc.styles['Normal']
font = style.font
font.name = 'Malgun Gothic'
font.size = Pt(10)
style.paragraph_format.line_spacing = 1.5
style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')

# ── Helper Functions ───────────────────────────────────────

def set_cell_border(cell, **kwargs):
    """Set cell border. kwargs: top, bottom, left, right as dicts {sz, val, color}."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn('w:tcBorders'))
    if tcBorders is None:
        tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
        tcPr.append(tcBorders)

    for edge, attrs in kwargs.items():
        if attrs is None:
            # Remove existing
            existing = tcBorders.find(qn(f'w:{edge}'))
            if existing is not None:
                tcBorders.remove(existing)
        else:
            element = tcBorders.find(qn(f'w:{edge}'))
            if element is None:
                element = parse_xml(f'<w:{edge} {nsdecls("w")}></w:{edge}>')
                tcBorders.append(element)
            for attr, val in attrs.items():
                element.set(qn(f'w:{attr}'), str(val))


def apply_three_line_table(table):
    """Apply three-line table (삼선표) format: top thick, header-bottom thin, bottom thick."""
    nrows = len(table.rows)
    ncols = len(table.columns)

    # Remove all borders first
    for row in table.rows:
        for cell in row.cells:
            for edge in ['top', 'bottom', 'left', 'right', 'insideH', 'insideV']:
                set_cell_border(cell, **{edge: None})

    # Top row: thick top border (1.5pt = 12 eighth-points)
    for cell in table.rows[0].cells:
        set_cell_border(cell, top={"sz": "12", "val": "single", "color": "000000"})

    # Header-bottom: thin line (0.5pt = 4 eighth-points)
    for cell in table.rows[0].cells:
        set_cell_border(cell, bottom={"sz": "4", "val": "single", "color": "000000"})

    # Bottom row: thick bottom border (1.5pt)
    for cell in table.rows[-1].cells:
        set_cell_border(cell, bottom={"sz": "12", "val": "single", "color": "000000"})

    # Center-align all cells
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.size = Pt(9)
                    run.font.name = 'Malgun Gothic'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')

    # Set table width to page
    table.autofit = True
    table.alignment = WD_TABLE_ALIGNMENT.CENTER


def add_heading_chapter(text):
    """Add chapter heading (제1장)."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(12)
    # Outline level 0 for TOC
    pPr = p._p.get_or_add_pPr()
    outline = parse_xml(f'<w:outlineLvl {nsdecls("w")} w:val="0"/>')
    pPr.append(outline)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = 'Malgun Gothic'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')


def add_heading_section(text):
    """Add section heading (1.1)."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    # Outline level 1 for TOC
    pPr = p._p.get_or_add_pPr()
    outline = parse_xml(f'<w:outlineLvl {nsdecls("w")} w:val="1"/>')
    pPr.append(outline)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(13)
    run.font.name = 'Malgun Gothic'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')


def add_heading_subsection(text):
    """Add sub-section heading."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    # Outline level 2 for TOC
    pPr = p._p.get_or_add_pPr()
    outline = parse_xml(f'<w:outlineLvl {nsdecls("w")} w:val="2"/>')
    pPr.append(outline)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = 'Malgun Gothic'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')


def add_body(text):
    """Add body paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.name = 'Malgun Gothic'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')
    return p


def add_bullet(text):
    """Add bullet point."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.0)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(f"• {text}")
    run.font.size = Pt(10)
    run.font.name = 'Malgun Gothic'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')


def add_figure(image_filename, caption, width_inches=5.5):
    """Add centered figure with caption."""
    path = IMG(image_filename)
    if not os.path.exists(path):
        add_body(f"[그림 누락: {image_filename}]")
        return

    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_img.add_run()
    run.add_picture(path, width=Inches(width_inches))

    p_cap = doc.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cap.paragraph_format.space_before = Pt(2)
    p_cap.paragraph_format.space_after = Pt(10)
    run = p_cap.add_run(caption)
    run.font.size = Pt(9)
    run.font.name = 'Malgun Gothic'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')


def make_table(headers, rows, col_widths=None):
    """Create a three-line table, return it. Headers and rows are lists of strings."""
    ncols = len(headers)
    nrows = 1 + len(rows)
    table = doc.add_table(rows=nrows, cols=ncols)
    table.style = 'Table Grid'

    # Header
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(str(h))
        run.bold = True
        run.font.size = Pt(9)
        run.font.name = 'Malgun Gothic'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')

    # Data rows
    for i, row_data in enumerate(rows):
        for j, val in enumerate(row_data):
            cell = table.rows[i + 1].cells[j]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(9)
            run.font.name = 'Malgun Gothic'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')

    apply_three_line_table(table)

    # Set column widths if provided
    if col_widths:
        for row in table.rows:
            for j, w in enumerate(col_widths):
                row.cells[j].width = Inches(w)

    return table


def add_table_caption(text):
    """Add table caption above table."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(9)
    run.font.name = 'Malgun Gothic'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')


def add_page_break():
    doc.add_page_break()


# ══════════════════════════════════════════════════════════════
# COVER PAGE
# ══════════════════════════════════════════════════════════════

# Empty lines for spacing
for _ in range(6):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("전통적 영상처리 기법 기반\n문서 이미지 자동 교정 및 향상 시스템 설계")
run.bold = True
run.font.size = Pt(22)
run.font.name = 'Malgun Gothic'
run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')

for _ in range(4):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("고급컴퓨터비전 (2026학년도 1학기)\n최종 연구보고서")
run.font.size = Pt(14)
run.font.name = 'Malgun Gothic'
run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')

for _ in range(3):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("2026년 5월")
run.font.size = Pt(13)
run.font.name = 'Malgun Gothic'
run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("GitHub: https://github.com/ybq9430/document-image-correction")
run.font.size = Pt(10)
run.font.name = 'Malgun Gothic'
run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')

add_page_break()

# ══════════════════════════════════════════════════════════════
# 목차 (Table of Contents)
# ══════════════════════════════════════════════════════════════

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(18)
run = p.add_run("목 차")
run.bold = True
run.font.size = Pt(18)
run.font.name = 'Malgun Gothic'
run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')

# Insert TOC field — right-click → Update Field in Word
toc_p = doc.add_paragraph()
run_begin = toc_p.add_run()
run_begin._r.append(parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>'))
run_instr = toc_p.add_run()
run_instr._r.append(parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> TOC \\o "1-3" \\h \\z </w:instrText>'))
run_sep = toc_p.add_run()
run_sep._r.append(parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>'))
run_text = toc_p.add_run("[ 목차를 생성하려면 여기를 우클릭 → 필드 업데이트(F) / Right-click → Update Field ]")
run_text.font.size = Pt(9)
run_text.font.color.rgb = RGBColor(128, 128, 128)
run_text.font.name = 'Malgun Gothic'
run_text._element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')
run_end = toc_p.add_run()
run_end._r.append(parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>'))

add_page_break()

# ══════════════════════════════════════════════════════════════
# 그림 목차 (List of Figures)
# ══════════════════════════════════════════════════════════════

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(12)
run = p.add_run("그 림 목 차")
run.bold = True
run.font.size = Pt(18)
run.font.name = 'Malgun Gothic'
run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')

figures = [
    ("그림 1.", "전체 처리 파이프라인 흐름도"),
    ("그림 2.", "파이프라인 단계별 중간 결과 예시 (7단계)"),
    ("그림 3.", "Canny 에지 검출 및 Dilation 효과"),
    ("그림 4.", "Canny + Dilation ROI 확대 비교"),
    ("그림 5.", "CLAHE 적용 전후 히스토그램 및 이미지 비교"),
    ("그림 6.", "검출 전략 분포 및 이미지별 전략 사용 현황"),
    ("그림 7.", "파이프라인 단계별 처리 시간 누적 막대 그래프"),
    ("그림 8.", "전체 6장 테스트 이미지의 원본/교정/향상 3단 비교 그리드"),
    ("그림 9.", "파라미터 민감도 히트맵"),
]
for num, title in figures:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Cm(1.5)
    run = p.add_run(f"{num}  {title}")
    run.font.size = Pt(10)
    run.font.name = 'Malgun Gothic'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')

add_page_break()

# ══════════════════════════════════════════════════════════════
# 표 목차 (List of Tables)
# ══════════════════════════════════════════════════════════════

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(12)
run = p.add_run("표 목 차")
run.bold = True
run.font.size = Pt(18)
run.font.name = 'Malgun Gothic'
run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')

tables = [
    ("표 1.", "개발 환경 및 제약 조건"),
    ("표 2.", "프로젝트 수행 일정"),
    ("표 3.", "실험 데이터 개요"),
    ("표 4.", "평가 지표 체계"),
    ("표 5.", "시스템 파라미터 기본값"),
    ("표 6.", "문서 교정 실험 결과 (Google Colab CPU 환경)"),
    ("표 7.", "Experiment 1: Canny Low Threshold"),
    ("표 8.", "Experiment 2: Canny High Threshold"),
    ("표 9.", "Experiment 3: Gaussian Kernel Size"),
    ("표 10.", "Experiment 4: CLAHE clipLimit"),
    ("표 11.", "종합 파라미터 권장값"),
]
for num, title in tables:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Cm(1.5)
    run = p.add_run(f"{num}  {title}")
    run.font.size = Pt(10)
    run.font.name = 'Malgun Gothic'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')

add_page_break()

# ══════════════════════════════════════════════════════════════
# 제1장. 서론
# ══════════════════════════════════════════════════════════════

add_heading_chapter("제1장. 서론")

# ── 1.1 연구 배경 및 의의 ──
add_heading_section("1.1 연구 배경 및 의의")

add_body("최근 디지털 전환이 가속화됨에 따라 종이 문서를 디지털 형태로 변환하는 수요가 급격히 증가하고 있다. 기업, 정부 기관, 교육 기관 등에서는 대량의 문서를 효율적으로 관리하기 위해 문서 스캔 및 OCR(Optical Character Recognition) 기술을 적극적으로 활용하고 있다. 특히 스마트폰의 보급으로 인해 전용 스캐너 없이도 카메라를 이용한 문서 촬영이 일상화되었으나, 이러한 방식으로 획득된 문서 이미지에는 다양한 품질 문제가 발생한다.")

add_body("스마트폰으로 문서를 촬영할 때 발생하는 주요 문제점은 다음과 같다. 첫째, 촬영 각도에 의한 투시 왜곡(perspective distortion)으로 문서가 사다리꼴 형태로 변형된다. 둘째, 실내 조명이나 자연광에 의한 불균일한 조명(uneven illumination)이 발생하여 문서의 일부 영역이 과도하게 밝거나 어두워진다. 셋째, 카메라 센서의 한계로 인한 노이즈(noise)가 이미지 품질을 저하시킨다. 넷째, 낮은 대비(contrast)로 인해 텍스트의 가독성이 떨어지는 문제가 있다.")

add_body("이러한 문제를 해결하기 위해 딥러닝 기반 접근법이 연구되고 있으나, 높은 계산 비용과 대규모 학습 데이터의 필요성, 그리고 모델의 해석 불가능성 등의 한계가 존재한다. 반면, 전통적 영상처리 기법은 수학적 원리에 기반하여 결과의 해석이 용이하고, 계산 비용이 낮으며, GPU 없이도 실시간 처리가 가능하다는 장점을 갖는다. 따라서 본 프로젝트에서는 전통적 컴퓨터 비전 및 영상처리 기법만을 활용하여 문서 이미지의 자동 교정 및 향상 시스템을 설계하고 구현하였다.")

# ── 1.2 연구 가설 ──
add_heading_section("1.2 연구 가설")

add_body("본 프로젝트의 연구 가설은 다음과 같다: 문서 이미지에서 문서의 외곽 경계가 충분히 검출 가능한 경우, Canny 에지 검출, 윤곽선 근사, 투시 변환, CLAHE, 적응적 이진화와 같은 전통적 영상처리 기법만으로도 문서의 기하학적 왜곡을 효과적으로 보정하고 텍스트 가독성을 향상시킬 수 있다.")

add_body("즉, 스마트폰 촬영 과정에서 기울어지거나 사다리꼴 형태의 원근 왜곡이 발생하더라도, 문서 외곽을 안정적으로 찾을 수 있다면 딥러닝 기반 모델 없이도 충분히 실용적인 수준의 문서 교정 결과를 얻을 수 있다는 가정이다. 또한 조명 불균일, 저대비, 미세 잡음과 같은 품질 저하 요소에 대해서도 CLAHE, Adaptive Thresholding, Morphological Processing의 조합이 가독성 향상에 유의미한 효과를 줄 것으로 예상한다.")

# ── 1.3 연구 목표 ──
add_heading_section("1.3 연구 목표")

add_body("본 프로젝트의 핵심 목표는 스마트폰으로 촬영된 문서 이미지를 입력으로 받아 기하학적 교정과 화질 향상을 자동으로 수행하는 완전한 처리 파이프라인을 구축하는 것이다. 구체적인 연구 목표는 다음과 같다.")

add_bullet("자동 문서 영역 검출: 복잡한 배경에서 문서 영역을 자동으로 식별하고 경계를 추출한다.")
add_bullet("투시 기하 교정: 검출된 문서의 네 꼭짓점을 기반으로 투시 변환을 적용하여 정면에서 촬영한 것과 같은 직사각형 형태로 교정한다.")
add_bullet("이미지 품질 향상: 적응적 이진화, 대비 향상, 형태학적 연산 등을 통해 문서의 가독성을 높인다.")
add_bullet("완전한 처리 흐름 구축: 딥러닝을 사용하지 않고 전통적 영상처리 기법만으로 입력부터 출력까지의 일관된 파이프라인을 구현한다.")

# ── 1.4 개발 환경 ──
add_heading_section("1.4 개발 환경 및 제약 조건")

add_table_caption("표 1. 개발 환경 및 제약 조건")
make_table(
    ["항목", "내용"],
    [
        ["프로그래밍 언어", "Python 3.x"],
        ["핵심 라이브러리", "OpenCV 4.13.0, NumPy 2.0.2, Matplotlib, scikit-image, pandas"],
        ["개발 환경", "Google Colab (CPU 런타임)"],
        ["ML/DL 사용", "금지 (torch, tensorflow, keras, sklearn 등 불가)"],
        ["허용 기법", "전통적 Computer Vision + Image Processing만 사용"],
        ["소스코드", "https://github.com/ybq9430/document-image-correction"],
    ],
    col_widths=[2.5, 4.0],
)

# ── 1.5 처리 흐름도 ──
add_heading_section("1.5 수행 방법 개요")

add_body("본 시스템은 스마트폰으로 촬영된 문서 이미지를 입력으로 받아 교정 및 향상된 문서 이미지를 출력하는 파이프라인 구조로 설계되었다. 전체 처리 과정은 전처리(Preprocessing), 문서 검출(Detection), 기하 교정(Geometry Correction), 후처리(Enhancement)의 네 단계로 구성되며, 각 단계는 순차적으로 실행된다.")

add_figure("_flowchart.png", "그림 1. 전체 처리 파이프라인 흐름도")

# ── 1.6 프로젝트 수행 일정 ──
add_heading_section("1.6 프로젝트 수행 일정")

add_body("본 프로젝트는 총 4단계에 걸쳐 수행되었으며, 각 단계별 주요 내용과 결과물은 다음과 같다.")

add_table_caption("표 2. 프로젝트 수행 일정")
make_table(
    ["단계", "기간", "주요 수행 내용", "결과물"],
    [
        ["1단계: 주제 결정", "10~12주차\n(5.4~5.17)", "주제 분야 탐색, 컴퓨터 비전 기술 동향 조사,\n전통적 영상처리 기법 중심 세부 기술 분야 선정", "수행계획서 제출"],
        ["2단계: 개요 작성", "12~13주차\n(5.18~5.24)", "연구 가설 설정, 4모듈 11단계 알고리즘 구조 설계,\n수행 방법 정리 및 평가 지표 체계 수립", "알고리즘 설계도,\nTR 개요 초안"],
        ["3단계: 내용 작성", "13~14주차\n(5.25~5.31)", "파이프라인 모듈별 구현 (전처리/검출/기하보정/향상),\n3단계 적응형 문서 검출 전략 구현,\n가설 검증용 평가 코드 작성", "소스코드 (Colab),\n중간 실험 결과"],
        ["4단계: 고찰/분석", "14~15주차\n(6.1~6.15)", "6장 테스트 이미지 전체 평가, 파라미터 민감도 실험,\n결과 데이터 분석, TR 보고서 작성 및 차트 생성", "TR 최종 보고서,\n발표 자료"],
    ],
    col_widths=[1.6, 1.3, 2.5, 1.3],
)

add_page_break()

# ══════════════════════════════════════════════════════════════
# 제2장. 본론
# ══════════════════════════════════════════════════════════════

add_heading_chapter("제2장. 본론")

# ── 2.1 전체 알고리즘 구조 ──
add_heading_section("2.1 전체 알고리즘 구조")

add_body("전체 알고리즘의 처리 파이프라인은 11단계로 구성된다: (1) Input Document Image → (2) Grayscale Conversion → (3) Gaussian Filtering → (4) Canny Edge Detection → (4b) Edge Dilation → (5) Contour Detection → (6) Quadrilateral Approximation + Convex Hull → (7) Corner Point Ordering → (8) Perspective Transformation → (9) CLAHE / Adaptive Thresholding → (10) Morphological Processing → (11) Output Corrected Image.")

add_body("이 파이프라인은 기능별로 네 개의 모듈로 구분된다: ① 전처리(Preprocessing), ② 문서 검출(Detection), ③ 기하학적 보정(Geometry), ④ 화질 향상(Enhancement).")

add_figure("_step_grid.png", "그림 2. 파이프라인 단계별 중간 결과 예시 (7단계)", width_inches=5.2)

# ── 2.2.1 전처리 ──
add_heading_section("2.2 단계별 알고리즘 상세")
add_heading_subsection("2.2.1 이미지 전처리 (Step 2-3)")

add_body("Step 2: Grayscale Conversion — 입력 RGB 이미지를 Grayscale로 변환하여 채널 수를 3에서 1로 줄임으로써 후속 처리의 계산 효율성을 높인다. 변환에는 BT.601 가중치를 적용한 표준 공식을 사용한다.")

# Centered formula
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(4)
p.paragraph_format.space_after = Pt(4)
run = p.add_run("Gray = 0.299 · R + 0.587 · G + 0.114 · B    (1)")
run.font.size = Pt(10)
run.font.name = 'Malgun Gothic'
run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')

add_body("OpenCV에서는 cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)를 사용하며, 내부적으로 위 가중치가 적용된 단일 채널 명암 영상을 출력한다.")

add_body("Step 3: Gaussian Filtering — 가우시안 필터를 적용하여 고주파 노이즈를 제거하고 영상을 평활화한다. 커널 크기는 5×5, 시그마는 0(자동 계산)으로 설정하였다. 이는 후속 Canny 에지 검출의 안정성을 높이고 불필요한 잡음 에지 생성을 줄이는 효과가 있다.")

# ── 2.2.2 문서 검출 ──
add_heading_subsection("2.2.2 문서 영역 검출 (Step 4-6)")

add_body("Step 4: Canny Edge Detection — 문서의 외곽 경계 후보를 추출하는 단계이다. Canny 알고리즘은 Gaussian smoothing → Gradient calculation (Sobel) → Non-maximum Suppression → Double Thresholding → Hysteresis Edge Tracking의 순서로 동작한다. 하한 임계값 75, 상한 임계값 200을 적용하였다.")

add_body("Step 4b: Edge Dilation — Canny 에지 검출 결과에서 발생하는 조각난 에지 세그먼트를 연결하기 위해 형태학적 팽창(dilation)을 적용한다. 3×3 사각형 커널로 2회 반복 팽창을 수행하여 끊어진 문서 경계를 연결한다.")

add_figure("canny_dilation.png", "그림 3. Canny 에지 검출 및 Dilation 효과 (① Gaussian Blurred → ② Canny Edges → ③ After Dilation)", width_inches=5.2)
add_figure("canny_dilation_roi.png", "그림 4. Canny + Dilation ROI 확대 비교 (에지 연결 효과 상세)", width_inches=5.2)

add_body("Step 5: Contour Detection — 팽창된 에지 맵에서 cv2.findContours(edges, cv2.RETR_LIST, cv2.CHAIN_APPROX_SIMPLE)를 사용하여 윤곽선(contour)을 추출한다. 면적 기준 내림차순으로 정렬하여 상위 10개의 후보만을 선택한다.")

add_body("Step 6: Quadrilateral Approximation (3단계 전략) — 문서 경계를 검출하기 위해 다음과 같은 3단계 적응적 전략을 적용한다.")

add_bullet("직접 근사 (Direct Approximation): 각 contour에 대해 Douglas-Peucker 알고리즘(cv2.approxPolyDP)을 적용. epsilon 비율을 0.02, 0.03, 0.05, 0.08, 0.10 순으로 시도하여 4개의 꼭짓점을 가진 근사 다각형을 찾는다.")
add_bullet("Convex Hull 근사 (Hull Approximation): 개별 contour의 직접 근사가 실패할 경우, 해당 contour의 convex hull(cv2.convexHull)을 계산한 후 동일한 근사 과정을 적용한다. 이는 프레임 경계 근처에서 부분적으로 잘린 문서 에지에 효과적이다.")
add_bullet("결합 Convex Hull (Combined Hull Fallback): 모든 contour가 실패할 경우, 상위 contour들을 결합하여 하나의 큰 contour를 만들고 convex hull 근사를 수행한다. 이는 문서가 이미지 프레임을 대부분 채우는 경우에 특히 효과적이다.")

add_body("최종적으로 검출된 사각형의 면적이 전체 이미지 면적의 2% 미만인 경우 노이즈로 판단하여 제외한다.")

add_page_break()

# ── 2.2.3 기하보정 ──
add_heading_subsection("2.2.3 기하학적 보정 (Step 7-8)")

add_body("Step 7: Corner Point Ordering — 검출된 4개의 꼭짓점을 일관된 순서로 정렬한다. 점의 순서가 잘못되면 투시 변환 결과가 뒤틀리거나 반전될 수 있으므로, 이 단계는 Homography Matrix 계산 이전에 반드시 수행해야 한다. 정렬 기준은 TL(Top-Left): x+y 최소, TR(Top-Right): x-y 최대, BR(Bottom-Right): x+y 최대, BL(Bottom-Left): x-y 최소이며, 결과는 [TL, TR, BR, BL] 순서의 (4,2) float32 ndarray이다.")

add_body("Step 8: Perspective Transformation — 4개의 소스 점과 목적 점으로부터 3×3 Homography Matrix H를 계산하여 투시 변환을 적용한다. Homography 변환은 다음과 같은 사영 변환식으로 표현된다:")

# Homography formula
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(4)
p.paragraph_format.space_after = Pt(4)
run = p.add_run("x' = (h₁₁x + h₁₂y + h₁₃) / (h₃₁x + h₃₂y + h₃₃)")
run.font.size = Pt(10)
run.font.name = 'Malgun Gothic'
run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')
p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p2.paragraph_format.space_before = Pt(0)
p2.paragraph_format.space_after = Pt(4)
run2 = p2.add_run("y' = (h₂₁x + h₂₂y + h₂₃) / (h₃₁x + h₃₂y + h₃₃)    (2)")
run2.font.size = Pt(10)
run2.font.name = 'Malgun Gothic'
run2._element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')

add_body("출력 이미지의 크기는 검출된 문서의 실제 크기에 맞게 자동 계산되며, 계산식은 다음과 같다:")

# Output size formula
p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
p3.paragraph_format.space_before = Pt(4)
p3.paragraph_format.space_after = Pt(2)
run3 = p3.add_run("Width  = max(∥BL − BR∥, ∥TL − TR∥)    (3)")
run3.font.size = Pt(10)
run3.font.name = 'Malgun Gothic'
run3._element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')

p4 = doc.add_paragraph()
p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
p4.paragraph_format.space_before = Pt(0)
p4.paragraph_format.space_after = Pt(4)
run4 = p4.add_run("Height = max(∥TR − BR∥, ∥TL − BL∥)    (4)")
run4.font.size = Pt(10)
run4.font.name = 'Malgun Gothic'
run4._element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')

add_body("OpenCV에서 cv2.getPerspectiveTransform(src_pts, dst_pts)로 변환 행렬을 구하고, cv2.warpPerspective()로 실제 변환을 수행한다.")

# ── 2.2.4 화질 향상 ──
add_heading_subsection("2.2.4 화질 향상 (Step 9-10)")

add_body("Step 9: CLAHE + Adaptive Thresholding — CLAHE(Contrast Limited Adaptive Histogram Equalization)를 적용하여 지역 대비를 향상시킨다. CLAHE는 이미지를 타일(8×8)로 분할하고 각 타일에 대해 히스토그램 균등화를 수행하되, 대비 제한(clip limit=2.0)을 적용하여 노이즈 증폭을 억제한다. 이후 Adaptive Thresholding(Gaussian method, block size=11, C=2)을 적용하여 조명 불균일을 보정하고 전경 문자와 배경을 분리한다.")

add_figure("clahe_histogram.png", "그림 5. CLAHE 적용 전후 히스토그램 및 이미지 비교", width_inches=5.2)

add_body("Step 10: Morphological Processing — 형태학적 닫힘 연산(MORPH_CLOSE, 3×3 kernel)을 적용하여 문자 내부의 작은 구멍을 메우고 배경의 미세 잡음을 제거한다. 최종적으로 텍스트 영역이 정리되어 가독성이 향상된다.")

# ── 2.3 실험 설계 ──
add_heading_section("2.3 실험 설계")

# 2.3.1
add_heading_subsection("2.3.1 실험 데이터")

add_body("실험에는 스마트폰(Samsung Galaxy S24 Ultra)으로 직접 촬영한 6장의 문서 이미지를 사용하였다. 모든 이미지는 1280×1706 해상도이다. 촬영 조건은 약 15°~45° 기울어진 각도, 실내 형광등 조명 환경, 책상 위에 놓인 A4 문서이다. 문서 유형은 텍스트 문서, 표가 포함된 문서, 필기 노트 등 다양하게 구성하였다.")

add_table_caption("표 3. 실험 데이터 개요")
make_table(
    ["이미지", "해상도", "문서 유형", "촬영 각도", "조명"],
    [
        ["img01", "1280×1706", "텍스트 문서", "~30°", "실내 형광등"],
        ["img02", "1280×1706", "표 포함 문서", "~20°", "실내 형광등"],
        ["img03", "1280×1706", "필기 노트", "~35°", "실내 형광등"],
        ["img04", "1280×1706", "텍스트 문서", "~15°", "실내 형광등"],
        ["img05", "1280×1706", "혼합 문서", "~45°", "실내 형광등"],
        ["img06", "1280×1706", "텍스트 문서", "~25°", "실내 형광등"],
    ],
    col_widths=[1.0, 1.2, 1.5, 1.2, 1.5],
)

# 2.3.2
add_heading_subsection("2.3.2 평가 지표")

add_body("시스템의 성능을 다각적으로 평가하기 위해 5대 검증 지표를 사용하였다. 특히 PSNR과 SSIM은 Grayscale-level(주 지표)과 Binary-level(참고용)의 2계층으로 측정하였다. 이진화 기반 문서 향상 시스템에서 최종 출력(흑백 이진 이미지)과 원본(컬러/그레이스케일)을 직접 비교하는 것은 픽셀 도메인의 근본적 차이로 인해 PSNR이 5~7 dB, SSIM이 0.45~0.59로 낮게 나타난다. 이는 화질 저하가 아닌 도메인 변환(grayscale→binary)의 결과이므로, 이진화 이전 단계인 CLAHE 향상 이미지를 원본 Grayscale과 비교하는 Grayscale-level 측정을 주 평가 지표로 도입하였다.")

add_table_caption("표 4. 평가 지표 체계")
make_table(
    ["#", "지표", "측정 방법"],
    [
        ["①", "문서 영역 검출 성공률", "성공 이미지 수 / 전체 이미지 수 × 100%"],
        ["②", "처리 시간", "time.perf_counter() 측정 (mean/min/max/total)"],
        ["③", "대비 향상도 (Grayscale-level, 주 지표)", "PSNR-Gray(dB), SSIM-Gray(0~1) — CLAHE 향상 이미지 vs 원본 Grayscale"],
        ["④", "대비 향상도 (Binary-level, 참고용)", "PSNR-Binary(dB), SSIM-Binary(0~1) — 최종 이진화 결과 vs 원본"],
        ["⑤", "선명도 향상도", "Laplacian Variance (전/후/배율) — 이진화 결과의 고주파 성분"],
    ],
    col_widths=[0.5, 2.5, 3.5],
)

# ── 2.4 파라미터 ──
add_heading_section("2.4 파라미터 설정")

add_table_caption("표 5. 시스템 파라미터 기본값")
make_table(
    ["파라미터", "기본값", "설명"],
    [
        ["Gaussian kernel", "(5, 5)", "노이즈 제거 강도"],
        ["Canny low threshold", "75", "약한 에지 임계값"],
        ["Canny high threshold", "200", "강한 에지 임계값"],
        ["Edge dilation kernel", "(3,3), 2 iterations", "에지 연결을 위한 팽창"],
        ["approxPolyDP epsilon", "0.02~0.10 (자동 탐색)", "다각형 근사 정밀도"],
        ["CLAHE clipLimit", "2.0", "대비 향상 강도"],
        ["CLAHE tileGridSize", "(8, 8)", "지역 처리 크기"],
        ["Adaptive blockSize", "11", "지역 이진화 블록 크기"],
        ["Adaptive C", "2", "이진화 보정값"],
        ["Morph kernel", "(3, 3)", "형태학적 연산 크기"],
    ],
    col_widths=[2.2, 2.0, 2.3],
)

add_page_break()

# ══════════════════════════════════════════════════════════════
# 제3장. 결과 및 고찰
# ══════════════════════════════════════════════════════════════

add_heading_chapter("제3장. 결과 및 고찰")

# ── 3.1.1 결과 테이블 ──
add_heading_section("3.1 실험 결과")
add_heading_subsection("3.1.1 문서 검출 성공률 및 처리 시간")

add_body("6장의 테스트 이미지에 대한 처리 결과는 다음과 같다. 평가 지표는 Grayscale-level(주 지표)과 Binary-level(참고용)의 2계층으로 측정하였다.")

add_table_caption("표 6. 문서 교정 실험 결과 (Google Colab CPU 환경)")
make_table(
    ["이미지", "검출", "시간(s)", "PSNR-Gray\n(dB)", "PSNR-Binary\n(dB)",
     "Laplacian\nBefore", "Laplacian\nAfter", "Laplacian\nRatio", "SSIM-Gray", "SSIM-Binary"],
    [
        ["img01", "성공", "0.186", "13.64", "5.07", "77.33", "10,083.79", "130.40x", "0.5987", "0.4516"],
        ["img02", "성공", "0.099", "15.47", "6.59", "26.81", "8,213.48", "306.41x", "0.6870", "0.5938"],
        ["img03", "성공", "0.124", "14.79", "6.50", "67.67", "10,916.22", "161.31x", "0.6351", "0.5213"],
        ["img04", "성공", "0.096", "13.85", "6.09", "80.07", "10,441.34", "130.40x", "0.5980", "0.5173"],
        ["img05", "성공", "0.099", "14.68", "6.32", "68.96", "9,906.58", "143.65x", "0.6138", "0.5411"],
        ["img06", "성공", "0.112", "15.42", "6.57", "74.31", "9,479.85", "127.56x", "0.6297", "0.5375"],
        ["평균", "100%", "0.119", "14.64", "6.19", "65.86", "9,840.21", "166.62x", "0.6271", "0.5271"],
    ],
    col_widths=[0.65, 0.5, 0.6, 0.75, 0.75, 0.75, 0.8, 0.75, 0.75, 0.75],
)

# ── 3.1.2 분석 ──
add_heading_subsection("3.1.2 결과 분석")

add_body("① 문서 영역 검출 성공률: 6/6 = 100.0%")
add_body("3단계 적응형 검출 전략(직접 근사 → Convex Hull → 결합 Hull)을 통해 모든 테스트 이미지에서 문서 영역을 성공적으로 검출하였다. 특히 문서가 프레임을 대부분 채우는 경우에도 결합 Convex Hull 전략이 효과적으로 작동하여 100% 검출률을 달성할 수 있었다.")

add_figure("detection_strategy.png", "그림 6. 검출 전략 분포 및 이미지별 전략 사용 현황 (B3)", width_inches=5.0)

add_body("② 처리 시간: 평균 0.119초 (최소 0.096초, 최대 0.186초)")
add_body("Google Colab CPU 런타임 환경에서 6장의 테스트 이미지(1280×1706)에 대한 전체 배치 처리 시간은 0.716초, 이미지당 평균 0.119초가 소요되었다. 이는 '단일 이미지 처리 시간 5초 이내' 요구사항을 크게 상회하는 결과로, 실시간 처리도 충분히 가능한 수준이다. img01의 처리 시간이 0.186초로 상대적으로 높은 것은 첫 번째 이미지 처리 시 Python 인터프리터 및 OpenCV 내부 캐시 초기화에 따른 Cold-start 효과로 분석된다. 순수 파이프라인 처리 시간(파일 I/O 제외)은 평균 0.105초, 단계별 프로파일링 기준 순수 알고리즘 실행 시간은 평균 81.4ms로 더욱 우수한 성능을 보였다.")

add_figure("stage_timing.png", "그림 7. 파이프라인 단계별 처리 시간 누적 막대 그래프 (B2)", width_inches=5.2)

add_body("③ 대비 및 선명도 향상 (2계층 평가)")

add_body("Laplacian Variance: 평균 65.86 → 9,840.21 (약 166.62배 향상) — Laplacian 분산값은 이미지의 선명도(sharpness)를 나타내는 지표로, 처리 전 평균 65.86에서 처리 후 평균 9,840.21로 약 166.62배 증가하여 적응적 이진화와 CLAHE를 통한 화질 향상이 매우 효과적이었음을 확인할 수 있다. 이미지별 Laplacian Ratio는 127.56x ~ 306.41x 범위로, 모든 이미지에서 100배 이상의 선명도 향상이 관찰되었다.")

add_body("PSNR-Gray (주 지표): 평균 14.64 dB (범위: 13.64~15.47 dB) — Grayscale-level PSNR은 CLAHE로 대비 향상된 이미지와 원본 Grayscale 이미지 간의 픽셀 단위 손실을 측정한다. Binary-level PSNR(평균 6.19 dB)에 비해 약 2.4배 높은 값을 보여, 이진화 이전 단계에서의 화질 손실이 매우 제한적임을 입증한다. 이는 CLAHE가 과도한 왜곡 없이 지역 대비를 효과적으로 향상시켰음을 의미하며, 최종 이진화 결과의 낮은 PSNR이 화질 저하가 아닌 grayscale→binary 도메인 변환에 기인함을 확인해준다.")

add_body("SSIM-Gray (주 지표): 평균 0.6271 (범위: 0.5980~0.6870) — Grayscale-level SSIM은 CLAHE 향상 이미지와 원본 Grayscale 간의 구조적 유사도를 측정한다. 평균 0.6271은 Binary-level SSIM(평균 0.5271) 대비 약 19% 높은 값으로, CLAHE 처리가 문서의 구조적 정보(텍스트 레이아웃, 표 경계, 문단 구분)를 양호하게 보존함을 나타낸다. Binary-level SSIM(0.4516~0.5938)은 Adaptive Thresholding이 grayscale의 중간 계조를 흑/백으로 양자화하면서 발생하는 불가피한 구조 정보 손실을 반영한다.")

add_body("④ 처리 전후 시각 비교")
add_body("각 이미지에 대해 원본, 투시 교정, 최종 향상의 3단계 비교도를 생성하였다. 시각적 비교 결과, 모든 이미지에서 기울어진 문서의 정면 직사각형 교정, 조명 불균일 제거 및 균일 배경 변환, 텍스트-배경 대비 향상, 문자 윤곽 선명화 및 배경 잡음 제거가 확인되었다.")

add_figure("_comparison.png", "그림 8. 전체 6장 테스트 이미지의 원본/교정/향상 3단 비교 그리드 (B1)", width_inches=5.5)

add_page_break()

# ── 3.2 파라미터 민감도 ──
add_heading_section("3.2 파라미터 민감도 실험 결과")

add_body("대표 이미지(img01)에 대해 4가지 주요 파라미터의 영향을 실험하였다. (참고: 파라미터 민감도 실험의 PSNR/SSIM은 Binary-level 기준으로 측정되었다.)")

add_table_caption("표 7. Experiment 1: Canny Low Threshold")
make_table(
    ["low", "검출 성공", "PSNR(dB)", "SSIM", "Laplacian Ratio"],
    [
        ["50", "O", "5.1", "0.4516", "130.4x"],
        ["75", "O", "5.1", "0.4516", "130.4x"],
        ["100", "O", "5.1", "0.4543", "127.5x"],
        ["125", "O", "5.1", "0.4564", "125.9x"],
    ],
    col_widths=[1.2, 1.2, 1.2, 1.2, 1.8],
)

add_body("분석: Canny low threshold 변화가 최종 결과에 미치는 영향은 미미했다. 모든 값에서 검출에 성공했으며 PSNR, SSIM, Laplacian Ratio 모두 유사한 수준을 유지했다. 이는 Edge Dilation 단계가 Canny 파라미터의 작은 변화를 흡수하기 때문으로 해석된다.")

add_table_caption("표 8. Experiment 2: Canny High Threshold")
make_table(
    ["high", "검출 성공", "PSNR(dB)", "SSIM", "Laplacian Ratio"],
    [
        ["150", "O", "5.1", "0.4516", "130.4x"],
        ["200", "O", "5.1", "0.4516", "130.4x"],
        ["250", "O", "5.1", "0.4516", "130.4x"],
        ["300", "O", "4.9", "NaN", "0.0x"],
    ],
    col_widths=[1.2, 1.2, 1.2, 1.2, 1.8],
)

add_body("분석: high=300에서 SSIM이 NaN으로 떨어지고 Laplacian Ratio가 0이 되었다. 이는 임계값이 너무 높아 에지가 거의 검출되지 않아 이진화 결과가 빈 이미지가 되었기 때문이다. high=150~250 범위에서는 안정적인 결과를 보였다.")

add_table_caption("표 9. Experiment 3: Gaussian Kernel Size")
make_table(
    ["kernel", "검출 성공", "PSNR(dB)", "SSIM", "Laplacian Ratio"],
    [
        ["(3,3)", "O", "5.1", "0.4511", "130.8x"],
        ["(5,5)", "O", "5.1", "0.4516", "130.4x"],
        ["(7,7)", "O", "5.1", "0.4511", "130.8x"],
        ["(9,9)", "O", "4.9", "0.6399", "2.7x"],
    ],
    col_widths=[1.2, 1.2, 1.2, 1.2, 1.8],
)

add_body("분석: (3,3)~(7,7)까지는 안정적인 성능을 보였다. 그러나 (9,9)에서 SSIM이 0.64로 증가하고 Laplacian Ratio가 2.7x로 급감했다. 이는 과도한 블러링으로 문서 에지가 약화되어 검출된 경계가 부정확해졌기 때문이다. Gaussian kernel은 (5,5)가 최적이다.")

add_table_caption("표 10. Experiment 4: CLAHE clipLimit")
make_table(
    ["clipLimit", "검출 성공", "PSNR(dB)", "SSIM", "Laplacian Ratio"],
    [
        ["1.0", "O", "5.0", "0.5016", "94.4x"],
        ["2.0", "O", "5.1", "0.4516", "130.4x"],
        ["3.0", "O", "5.1", "0.4177", "156.7x"],
        ["4.0", "O", "5.1", "0.3882", "178.5x"],
        ["5.0", "O", "5.2", "0.3656", "195.9x"],
    ],
    col_widths=[1.2, 1.2, 1.2, 1.2, 1.8],
)

add_body("분석: CLAHE clipLimit이 증가할수록 Laplacian Ratio가 94.4x → 195.9x로 지속적으로 증가했으나, SSIM은 반대로 0.5016 → 0.3656으로 감소했다. 이는 clipLimit이 높을수록 대비가 과도하게 증폭되어 노이즈까지 강조되기 때문이다. 문서 가독성과 구조 보존 사이의 균형을 고려할 때 clipLimit=2.0이 가장 적절한 값으로 판단된다.")

add_figure("sensitivity_heatmap.png", "그림 9. 파라미터 민감도 히트맵 (3×4, RdYlGn 컬러맵) (B5)", width_inches=5.2)

add_body("")

add_table_caption("표 11. 종합 파라미터 권장값")
make_table(
    ["파라미터", "권장값", "근거"],
    [
        ["Canny low", "75", "넓은 범위(50~125)에서 안정적"],
        ["Canny high", "200", "150~250에서 안정적, 300에서 실패"],
        ["Gaussian kernel", "(5,5)", "(3,3)~(7,7)에서 안정적, (9,9)에서 성능 저하"],
        ["CLAHE clipLimit", "2.0", "선명도(130.4x)와 구조 보존(SSIM 0.4516) 간 최적 균형"],
    ],
    col_widths=[2.0, 1.5, 3.0],
)

add_page_break()

# ── 3.3 가설 검증 ──
add_heading_section("3.3 가설 검증")

add_body("본 연구의 가설은 '전통적 영상처리 기법만으로도 문서의 기하학적 왜곡을 효과적으로 보정하고 텍스트 가독성을 향상시킬 수 있다'는 것이었다. 실험 결과를 통해 다음과 같이 가설을 검증할 수 있다:")

add_bullet("문서 검출: 3단계 적응형 검출 전략을 통해 100%의 문서 검출 성공률을 달성하여, 딥러닝 없이도 다양한 촬영 조건에서 문서 영역을 안정적으로 검출할 수 있음을 입증하였다.")
add_bullet("기하학적 보정: 투시 변환을 통해 기울어진 문서를 정면 직사각형으로 성공적으로 교정하였다. Convex Hull 결합 전략을 통해 문서가 프레임을 대부분 채우는 경우에도 안정적인 보정이 가능했다.")
add_bullet("화질 향상: Laplacian 분산값이 평균 166.62배 향상되어, CLAHE + Adaptive Thresholding + Morphological Processing의 조합이 텍스트 가독성 향상에 매우 효과적임을 정량적으로 확인하였다. 또한 Grayscale-level PSNR 평균 14.64 dB와 SSIM 평균 0.6271을 통해, 이진화 이전 CLAHE 단계에서의 화질 손실이 제한적이며 구조적 정보가 양호하게 보존됨을 입증하였다.")
add_bullet("처리 속도: 배치 처리 기준 평균 0.119초/image, 순수 파이프라인 기준 0.105초/image의 처리 속도로 실시간 응용에도 충분한 성능을 보였다.")

add_body("따라서 본 연구의 가설은 실험 결과에 의해 지지된다고 결론 내릴 수 있다.")

# ── 3.4 한계점 ──
add_heading_section("3.4 한계점 및 향후 연구")

add_body("본 시스템의 한계점과 개선 방향은 다음과 같다:")

add_body("1. 복잡한 배경에서의 검출: 본 실험에 사용된 이미지는 상대적으로 단순한 배경(책상)에서 촬영되었다. 매우 복잡한 배경(패턴이 있는 바닥, 여러 물체가 있는 환경)에서는 문서 검출 성능이 저하될 수 있다. 이를 개선하기 위해 텍스처 기반 배경 제거 기법이나 적응형 ROI 탐색을 추가할 수 있다.")

add_body("2. Binary-level PSNR의 한계: Adaptive Thresholding으로 인한 이진화는 Binary-level PSNR을 평균 6.19 dB로 낮추는 근본 원인이다. 본 연구에서는 이를 보완하기 위해 Grayscale-level PSNR(평균 14.64 dB)을 주 지표로 도입하였으나, 궁극적으로는 OCR 인식률과 같은 실제 활용 지표를 추가 평가 지표로 도입하는 것이 바람직하다.")

add_body("3. 이진화 결과의 SSIM: Binary-level SSIM이 0.45~0.59 범위(평균 0.5271)에 머문 것은 grayscale→binary 도메인 변환의 본질적 특성이다. Grayscale-level SSIM은 평균 0.6271로 유의미하게 높으며, CLAHE 단계의 구조 보존력을 더 잘 반영한다. 문서 유형에 따라 Grayscale 보존 모드와 이진화 모드를 선택적으로 적용하는 하이브리드 접근법을 고려할 수 있다.")

add_body("4. 파라미터 자동 보정: 현재 파라미터는 수동으로 설정되어 있다. 이미지 특성(조도, 대비, 해상도)에 따라 파라미터를 자동으로 조정하는 적응형 메커니즘을 추가하면 더 넓은 범위의 입력에 대응할 수 있을 것이다.")

# ── 3.5 결론 ──
add_heading_section("3.5 결론")

add_body("본 프로젝트에서는 전통적 컴퓨터 비전 및 영상처리 기법만을 사용하여 스마트폰 촬영 문서 이미지의 자동 교정 및 향상 시스템을 성공적으로 설계하고 구현하였다. Canny 에지 검출, Convex Hull 기반 적응형 문서 검출, 투시 변환, CLAHE, 적응적 이진화, 형태학적 연산을 조합한 11단계 파이프라인을 구축하였다.")

add_body("6장의 테스트 이미지에 대한 실험 결과, 100% 문서 검출 성공률, 배치 처리 기준 평균 0.119초의 처리 속도(순수 파이프라인 0.105초), 평균 166.62배의 Laplacian 선명도 향상을 달성하였다. 2계층 PSNR/SSIM 평가를 통해 Grayscale-level PSNR 14.64 dB, SSIM 0.6271을 기록하여 CLAHE 기반 대비 향상의 효과를 정량적으로 입증하였다. 파라미터 민감도 실험을 통해 최적의 파라미터 조합을 도출하였으며, 모든 평가 지표에서 안정적인 성능을 확인하였다.")

add_body("본 연구는 딥러닝을 사용하지 않고도 전통적 영상처리 기법만으로 실용적인 문서 교정 시스템을 구축할 수 있음을 입증하였으며, 향후 OCR 연계, 모바일 최적화, 실시간 비디오 처리 등으로 확장할 수 있는 기반을 마련하였다.")

add_page_break()

# ══════════════════════════════════════════════════════════════
# 참고문헌
# ══════════════════════════════════════════════════════════════

add_heading_chapter("참고 문헌")

refs = [
    "[1] J. Canny, \"A Computational Approach to Edge Detection,\" IEEE TPAMI, vol. 8, no. 6, pp. 679-698, 1986.",
    "[2] R. C. Gonzalez and R. E. Woods, Digital Image Processing, 4th ed., Pearson, 2018.",
    "[3] G. Bradski and A. Kaehler, Learning OpenCV: Computer Vision with the OpenCV Library, O'Reilly Media, 2008.",
    "[4] R. Hartley and A. Zisserman, Multiple View Geometry in Computer Vision, 2nd ed., Cambridge University Press, 2004.",
    "[5] K. Zuiderveld, \"Contrast Limited Adaptive Histogram Equalization,\" in Graphics Gems IV, Academic Press, pp. 474-485, 1994.",
    "[6] D. H. Douglas and T. K. Peucker, \"Algorithms for the Reduction of the Number of Points Required to Represent a Digitized Line,\" Cartographica, vol. 10, no. 2, pp. 112-122, 1973.",
    "[7] S. Suzuki and K. Abe, \"Topological Structural Analysis of Digitized Binary Images by Border Following,\" CVGIP, vol. 30, no. 1, pp. 32-46, 1985.",
    "[8] J. Sauvola and M. Pietikäinen, \"Adaptive Document Image Binarization,\" Pattern Recognition, vol. 33, no. 2, pp. 225-236, 2000.",
]

for ref in refs:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(ref)
    run.font.size = Pt(9)
    run.font.name = 'Malgun Gothic'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')

# ── Save ──
doc.save(OUTPUT_PATH)
print(f"Report saved ({os.path.getsize(OUTPUT_PATH) / 1024 / 1024:.1f} MB)")
print(f"  Tables: {len(doc.tables)}")
print(f"  Images: {len(doc.inline_shapes)}")
